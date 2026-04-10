"""Report screen — AI-drafted report with 7 sections + manual editing.

Each section can be:
  - Generated via Sonnet (streamed to UI)
  - Manually edited after generation
  - Re-generated if edits are unsatisfactory
  - Exported as DOCX
"""

from __future__ import annotations
import threading, tkinter as tk
import customtkinter as ctk
from bc4d_intel import constants as C
from bc4d_intel.ui import widgets as W
from bc4d_intel.ai.prompts import REPORT_SECTIONS

SECTION_LABELS = {
    "executive_summary": "1. Executive Summary",
    "method_sample": "2. Method & Sample",
    "quantitative_results": "3. Quantitative Results",
    "qualitative_findings": "4. Qualitative Findings",
    "pre_post_comparison": "5. Pre/Post Comparison",
    "recommendations": "6. Recommendations",
    "appendix": "7. Appendix",
}


class ReportScreen(ctk.CTkFrame):
    def __init__(self, master, app):
        super().__init__(master, fg_color=C.BG, corner_radius=0)
        self.app = app
        self._sections = {}  # {section_name: text}
        self._current_section = None
        self._build()

    def _build(self):
        from bc4d_intel.ui.guide import workflow_steps

        # ── Top bar ──
        top = W.make_toolbar(self, height=80)
        inner = ctk.CTkFrame(top, fg_color="transparent")
        inner.pack(fill="x", padx=30, pady=8)

        row = ctk.CTkFrame(inner, fg_color="transparent")
        row.pack(fill="x")

        W.heading(row, "Report", size=22).pack(side="left")

        # Buttons — right-aligned, generous sizing
        self._status_lbl = ctk.CTkLabel(row, text="",
                                         font=ctk.CTkFont(family="Segoe UI", size=11),
                                         text_color=C.MUTED)
        self._status_lbl.pack(side="right", padx=(8, 0))

        ctk.CTkButton(
            row, text="Export DOCX", width=120, height=36,
            fg_color=C.BTN, hover_color=C.SELECT, text_color=C.TEXT,
            font=ctk.CTkFont(family="Segoe UI", size=12),
            corner_radius=6, command=self._export_docx,
        ).pack(side="right", padx=(8, 0))

        self._generate_all_btn = ctk.CTkButton(
            row, text="Generate Full Report", width=200, height=36,
            fg_color=C.ACCENT, hover_color=C.SELECT,
            font=ctk.CTkFont(family="Segoe UI", size=12, weight="bold"),
            corner_radius=6, command=self._generate_all,
        )
        self._generate_all_btn.pack(side="right", padx=(8, 0))

        ctk.CTkLabel(inner,
            text="Generate all sections with AI, or select a section on the left to generate individually.",
            font=ctk.CTkFont(family="Segoe UI", size=9),
            text_color=C.MUTED).pack(anchor="w", pady=(2, 0))

        # ── Body: section list (left) + editor (right) ──
        body = ctk.CTkFrame(self, fg_color="transparent")
        body.pack(fill="both", expand=True, padx=20, pady=10)
        body.columnconfigure(0, weight=1)
        body.columnconfigure(1, weight=3)
        body.rowconfigure(0, weight=1)

        # Left: section list
        left = W.make_card(body)
        left.grid(row=0, column=0, sticky="nsew", padx=(0, 5))

        ctk.CTkLabel(left, text="Report Sections",
                     font=ctk.CTkFont(family="Segoe UI", size=12, weight="bold"),
                     text_color=C.TEXT).pack(anchor="w", padx=12, pady=(10, 4))

        self._section_frame = ctk.CTkFrame(left, fg_color="transparent")
        self._section_frame.pack(fill="both", expand=True, padx=6, pady=6)

        for key, label in SECTION_LABELS.items():
            btn = ctk.CTkButton(
                self._section_frame,
                text=label, anchor="w", height=36,
                fg_color=C.PANEL, hover_color=C.SELECT,
                text_color=C.TEXT,
                font=ctk.CTkFont(family="Segoe UI", size=11),
                corner_radius=6,
                command=lambda k=key: self._select_section(k),
            )
            btn.pack(fill="x", padx=4, pady=2)

        # Right: scrollable content area (text + charts + tables)
        right = ctk.CTkFrame(body, fg_color="transparent")
        right.grid(row=0, column=1, sticky="nsew", padx=(5, 0))

        editor_header = ctk.CTkFrame(right, fg_color="transparent")
        editor_header.pack(fill="x", padx=4, pady=(4, 4))

        self._section_title = ctk.CTkLabel(
            editor_header, text="Select a section to edit",
            font=ctk.CTkFont(family="Segoe UI", size=13, weight="bold"),
            text_color=C.TEXT)
        self._section_title.pack(side="left")

        self._gen_btn = ctk.CTkButton(
            editor_header, text="Generate This Section", width=170, height=32,
            fg_color=C.ACCENT, hover_color=C.SELECT,
            font=ctk.CTkFont(family="Segoe UI", size=11, weight="bold"),
            corner_radius=6,
            command=self._generate_current)
        self._gen_btn.pack(side="right")

        # Scrollable content area
        self._content_scroll = ctk.CTkScrollableFrame(right, fg_color="transparent")
        self._content_scroll.pack(fill="both", expand=True)

        # Text editor
        text_container = ctk.CTkFrame(self._content_scroll, fg_color=C.ENTRY_BG,
                                       corner_radius=6, height=300)
        text_container.pack(fill="x", padx=4, pady=(0, 6))
        text_container.pack_propagate(False)

        self._editor = tk.Text(
            text_container, wrap="word", relief="flat", borderwidth=0,
            bg=C.ENTRY_BG, fg=C.TEXT, insertbackground=C.TEXT,
            font=("Segoe UI", 11), padx=14, pady=10,
            selectbackground=C.SELECT, selectforeground=C.TEXT,
            undo=True, spacing1=2, spacing3=2,
        )
        scrollbar = tk.Scrollbar(text_container, command=self._editor.yview)
        self._editor.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side="right", fill="y")
        self._editor.pack(side="left", fill="both", expand=True)

        # Chart/table area below editor
        self._viz_frame = ctk.CTkFrame(self._content_scroll, fg_color="transparent")
        self._viz_frame.pack(fill="x", padx=4, pady=(0, 6))

        # Markdown-style text tags for readable formatting
        self._editor.tag_configure("h2", font=("Segoe UI", 15, "bold"),
                                    spacing1=12, spacing3=4, foreground=C.ACCENT)
        self._editor.tag_configure("h3", font=("Segoe UI", 12, "bold"),
                                    spacing1=8, spacing3=2)
        self._editor.tag_configure("bold", font=("Segoe UI", 11, "bold"))
        self._editor.tag_configure("bullet", lmargin1=20, lmargin2=32,
                                    font=("Segoe UI", 11))

    def _select_section(self, section_key):
        """Load a section into the editor with markdown-style formatting + charts."""
        self._current_section = section_key
        label = SECTION_LABELS.get(section_key, section_key)
        self._section_title.configure(text=label)

        # Clear viz area — close any matplotlib figures to prevent memory leak
        try:
            import matplotlib.pyplot as plt
            for w in self._viz_frame.winfo_children():
                if hasattr(w, '_canvas'):
                    plt.close('all')
                    break
        except Exception:
            pass
        for w in self._viz_frame.winfo_children():
            w.destroy()

        # Show relevant chart/table for this section
        self._render_section_viz(section_key)

        self._editor.delete("1.0", "end")
        text = self._sections.get(section_key, "")
        if text:
            self._render_markdown(text)
        else:
            self._editor.insert("1.0", f"Click 'Generate' to create this section via AI.\n\n"
                                        f"Or type your own content here.")

    def _render_markdown(self, text: str):
        """Insert text with markdown-style formatting (headings, bold, bullets)."""
        import re
        self._editor.delete("1.0", "end")
        for line in text.split("\n"):
            stripped = line.strip()
            if stripped.startswith("### "):
                self._editor.insert("end", stripped[4:] + "\n", "h3")
            elif stripped.startswith("## "):
                self._editor.insert("end", stripped[3:] + "\n", "h2")
            elif stripped.startswith("- ") or stripped.startswith("* "):
                self._editor.insert("end", "\u2022 " + stripped[2:] + "\n", "bullet")
            elif "**" in stripped:
                # Parse bold markers
                parts = re.split(r'(\*\*[^*]+\*\*)', stripped)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        self._editor.insert("end", part[2:-2], "bold")
                    else:
                        self._editor.insert("end", part)
                self._editor.insert("end", "\n")
            else:
                self._editor.insert("end", line + "\n")

    def _generate_current(self):
        """Generate the current section via Sonnet."""
        if not self._current_section:
            return
        api_key = self.app.app_state.api_key
        if not api_key:
            self._status_lbl.configure(text="Set API key in Settings.", text_color=C.DANGER)
            return

        section = self._current_section
        self._gen_btn.configure(state="disabled", text="Generating...")
        self._editor.delete("1.0", "end")
        self._editor.insert("1.0", "Generating...\n")

        def work():
            from bc4d_intel.ai.report_writer import generate_section

            context = self._build_context()

            def on_stream(chunk):
                self.after(0, lambda c=chunk: self._append_text(c))

            # Clear editor before streaming
            self.after(0, lambda: self._editor.delete("1.0", "end"))

            text = generate_section(section, context, api_key, stream_cb=on_stream)
            self._sections[section] = text

            # Save to app state
            self.app.app_state.report_sections[section] = text
            self.app.app_state.save()

            # Re-render with formatting after streaming completes
            self.after(0, lambda t=text: self._render_markdown(t))
            self.after(0, lambda: self._gen_btn.configure(state="normal", text="Generate"))
            self.after(0, lambda: self._status_lbl.configure(
                text=f"Section generated.", text_color=C.SUCCESS))

        threading.Thread(target=work, daemon=True).start()

    def _append_text(self, chunk):
        """Append streamed text to editor."""
        self._editor.insert("end", chunk)
        self._editor.see("end")

    def _generate_all(self):
        """Generate all 7 sections sequentially."""
        api_key = self.app.app_state.api_key
        if not api_key:
            self._status_lbl.configure(text="Set API key in Settings.", text_color=C.DANGER)
            return

        self._generate_all_btn.configure(state="disabled", text="Generating...")

        def work():
            from bc4d_intel.ai.report_writer import generate_section, build_data_context

            context = self._build_context()

            generated_so_far = {}
            for i, (section, label) in enumerate(SECTION_LABELS.items()):
                self.after(0, lambda l=label, n=i+1: self._status_lbl.configure(
                    text=f"Generating {n}/7: {l}...", text_color=C.ACCENT))

                text = generate_section(section, context, api_key,
                                        previous_sections=generated_so_far)
                self._sections[section] = text
                self.app.app_state.report_sections[section] = text
                generated_so_far[section] = text

            self.app.app_state.save()

            self.after(0, lambda: self._generate_all_btn.configure(
                state="normal", text="Generate All Sections"))
            self.after(0, lambda: self._status_lbl.configure(
                text="All 7 sections generated.", text_color=C.SUCCESS))
            # Show first section
            self.after(0, lambda: self._select_section("executive_summary"))

        threading.Thread(target=work, daemon=True).start()

    def _build_context(self) -> str:
        """Build rich data context with actual statistics for Sonnet."""
        from bc4d_intel.ai.report_writer import build_data_context
        match_result = getattr(self.app, "_match_result", None)
        import_screen = self.app._frames.get("import")

        pre_roles = import_screen._pre_roles if import_screen else None
        post_roles = import_screen._post_roles if import_screen else None
        pre_df = import_screen._pre_df if import_screen else None
        post_df = import_screen._post_df if import_screen else None

        # Use qualitative data from AI Analysis results (not old validation screen)
        tagged_data = None
        if self.app.app_state.tagged_responses:
            tagged_data = {}
            for label, classifications in self.app.app_state.tagged_responses.items():
                tagged_data[label] = [
                    {"text": c.get("text", ""),
                     "tag": c.get("cluster_title", ""),
                     "human_override": c.get("human_override", "")}
                    for c in classifications
                ]

        return build_data_context(
            self.app.app_state, match_result, tagged_data,
            pre_roles=pre_roles, post_roles=post_roles,
            pre_df=pre_df, post_df=post_df,
        )

    # ── Section visualizations (tables + charts) ────────────────

    def _render_section_viz(self, section_key):
        """Render relevant tables and charts below the text for this section."""
        try:
            if section_key == "executive_summary":
                self._viz_summary_table()
            elif section_key == "quantitative_results":
                self._viz_likert_table()
                self._viz_likert_chart()
            elif section_key == "qualitative_findings":
                self._viz_qualitative_tables()
            elif section_key == "pre_post_comparison":
                self._viz_comparison_table()
                self._viz_comparison_chart()
            elif section_key == "appendix":
                self._viz_full_stats_table()
        except Exception as e:
            ctk.CTkLabel(self._viz_frame, text=f"(Visualization not available: {e})",
                         font=ctk.CTkFont(size=9), text_color=C.MUTED
                         ).pack(anchor="w", padx=8)

    def _make_table(self, parent, headers, rows, title=""):
        """Build a formatted table using CTkFrame rows."""
        if title:
            ctk.CTkLabel(parent, text=title,
                         font=ctk.CTkFont(family="Segoe UI", size=12, weight="bold"),
                         text_color=C.TEXT).pack(anchor="w", padx=8, pady=(8, 4))

        table = ctk.CTkFrame(parent, fg_color=C.PANEL, corner_radius=6)
        table.pack(fill="x", padx=8, pady=(0, 8))

        # Header row
        hrow = ctk.CTkFrame(table, fg_color=C.ACCENT, corner_radius=4)
        hrow.pack(fill="x", padx=2, pady=(2, 0))
        for j, h in enumerate(headers):
            ctk.CTkLabel(hrow, text=h, width=max(60, 120 if j == 0 else 70),
                         font=ctk.CTkFont(family="Segoe UI", size=9, weight="bold"),
                         text_color="#ffffff", anchor="w"
                         ).pack(side="left", padx=4, pady=3)

        # Data rows
        for i, row_data in enumerate(rows):
            bg = C.CARD if i % 2 == 0 else C.PANEL
            drow = ctk.CTkFrame(table, fg_color=bg)
            drow.pack(fill="x", padx=2)
            for j, val in enumerate(row_data):
                ctk.CTkLabel(drow, text=str(val),
                             width=max(60, 120 if j == 0 else 70),
                             font=ctk.CTkFont(family="Consolas", size=9),
                             text_color=C.TEXT, anchor="w"
                             ).pack(side="left", padx=4, pady=2)

    def _get_import_data(self):
        """Get import screen data (dataframes + roles)."""
        imp = self.app._frames.get("import")
        if not imp:
            return None, None, None, None
        return (getattr(imp, "_pre_df", None), getattr(imp, "_pre_roles", None),
                getattr(imp, "_post_df", None), getattr(imp, "_post_roles", None))

    def _viz_summary_table(self):
        """Executive summary: key metrics table."""
        s = self.app.app_state
        rows = [
            ["Pre-Survey (Vorabfragebogen)", str(s.n_pre)],
            ["Post-Survey (Nachbefragung)", str(s.n_post)],
            ["Gematchte Paare", str(s.matched_pairs)],
            ["Match-Rate", f"{s.matched_pairs / max(s.n_post, 1) * 100:.0f}%"],
            ["Dropout (nur Pre)", str(s.unmatched_pre)],
        ]
        self._make_table(self._viz_frame, ["Kennzahl", "Wert"], rows,
                         title="Stichprobe")

    def _viz_likert_table(self):
        """Quantitative results: Likert items table."""
        pre_df, pre_roles, post_df, post_roles = self._get_import_data()
        if post_df is None or post_roles is None:
            return
        from bc4d_intel.core.stats_engine import analyze_all_likert
        items = analyze_all_likert(post_df, post_roles)
        rows = []
        for item in items:
            s = item["stats"]
            if s["mean"] is not None:
                rows.append([
                    item["label"][:45],
                    str(s["n"]),
                    str(s["mean"]),
                    str(s["sd"]),
                    f"{s['pct_agree']}%",
                ])
        if rows:
            self._make_table(self._viz_frame,
                             ["Item", "N", "M", "SD", "Zust."],
                             rows, title="Nachbefragung: Likert-Skalen")

    def _viz_likert_chart(self):
        """Quantitative results: horizontal bar chart."""
        pre_df, pre_roles, post_df, post_roles = self._get_import_data()
        if post_df is None or post_roles is None:
            return
        try:
            from bc4d_intel.core.chart_builder import _ensure_mpl, _apply_style
            _ensure_mpl()
            import matplotlib.pyplot as plt
            from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
            from bc4d_intel.core.stats_engine import analyze_all_likert

            items = analyze_all_likert(post_df, post_roles)
            items = [i for i in items if i["stats"]["mean"] is not None]
            if not items:
                return

            labels = [i["label"][:35] for i in items]
            means = [i["stats"]["mean"] for i in items]
            n = len(items)

            from bc4d_intel.core.chart_builder import _chart_colors
            cc = _chart_colors()

            fig, ax = plt.subplots(figsize=(8, max(2.5, n * 0.4)))
            _apply_style(fig, ax)
            y = range(n)
            colors = ["#059669" if m >= 3.5 else "#d97706" if m >= 2.5 else "#dc2626"
                      for m in means]
            ax.barh(y, means, height=0.6, color=colors, edgecolor=cc["edge"])
            ax.set_yticks(y)
            ax.set_yticklabels(labels, fontsize=9)
            ax.set_xlabel("Mittelwert (1-5)", fontsize=10)
            ax.set_xlim(1, 5)
            ax.invert_yaxis()
            ax.axvline(x=3, color=cc["muted"], linewidth=0.5, linestyle="--")
            for i, m in enumerate(means):
                ax.text(m + 0.05, i, f"{m}", va="center", fontsize=9, color=cc["text"])
            fig.tight_layout()

            card = W.make_card(self._viz_frame)
            card.pack(fill="x", padx=8, pady=4)
            canvas = FigureCanvasTkAgg(fig, master=card)
            canvas.draw()
            canvas.get_tk_widget().pack(fill="x", padx=4, pady=4)
        except Exception:
            pass

    def _viz_qualitative_tables(self):
        """Qualitative findings: category distribution per question."""
        results = getattr(self.app, "_analysis_results", {})
        if not results:
            tagged = self.app.app_state.tagged_responses
            if not tagged:
                return
            results = {k: {"classifications": v} for k, v in tagged.items()}

        for label, data in list(results.items())[:5]:  # cap at 5 questions
            classifications = data.get("classifications", [])
            if not classifications:
                continue
            from collections import Counter
            total = len(classifications)
            counts = Counter(
                c.get("cluster_title", "?") for c in classifications)
            rows = [[cat[:35], str(n), f"{n/total*100:.0f}%"]
                    for cat, n in counts.most_common(8)]
            short = label.split("]")[-1].strip()[:40] if "]" in label else label[:40]
            self._make_table(self._viz_frame,
                             ["Kategorie", "N", "%"], rows,
                             title=short)

    def _viz_comparison_table(self):
        """Pre/Post comparison: matched panel stats table."""
        match_result = getattr(self.app, "_match_result", None)
        pre_df, pre_roles, post_df, post_roles = self._get_import_data()
        if not match_result or not pre_roles or not post_roles:
            return
        matched = match_result.get("matched")
        if matched is None or len(matched) == 0:
            return
        from bc4d_intel.core.stats_engine import analyze_matched_likert
        comparisons = analyze_matched_likert(matched, pre_roles, post_roles)
        rows = []
        for c in comparisons:
            comp = c.get("comparison", {})
            if "error" in comp:
                continue
            sig = "*" if comp.get("significant") else ""
            rows.append([
                c["label"][:35],
                str(comp["pre_mean"]),
                str(comp["post_mean"]),
                f"{'+' if comp['mean_change'] > 0 else ''}{comp['mean_change']}",
                str(comp["cohens_d"]),
                f"{comp['effect_label']}{sig}",
            ])
        if rows:
            self._make_table(self._viz_frame,
                             ["Item", "Pre", "Post", "Diff", "d", "Effekt"],
                             rows, title="Pre/Post-Vergleich (gematchtes Panel)")

    def _viz_comparison_chart(self):
        """Pre/Post comparison: grouped bar chart."""
        match_result = getattr(self.app, "_match_result", None)
        pre_df, pre_roles, post_df, post_roles = self._get_import_data()
        if not match_result or not pre_roles or not post_roles:
            return
        matched = match_result.get("matched")
        if matched is None or len(matched) == 0:
            return
        try:
            from bc4d_intel.core.chart_builder import _ensure_mpl, _apply_style
            _ensure_mpl()
            import matplotlib.pyplot as plt
            from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
            import numpy as np
            from bc4d_intel.core.stats_engine import analyze_matched_likert

            comparisons = analyze_matched_likert(matched, pre_roles, post_roles)
            comps = [c for c in comparisons if "error" not in c.get("comparison", {})]
            if not comps:
                return

            labels = [c["label"][:30] for c in comps]
            pre_means = [c["comparison"]["pre_mean"] for c in comps]
            post_means = [c["comparison"]["post_mean"] for c in comps]
            n = len(comps)

            from bc4d_intel.core.chart_builder import _chart_colors
            cc = _chart_colors()

            fig, ax = plt.subplots(figsize=(8, max(2.5, n * 0.5)))
            _apply_style(fig, ax)
            y = np.arange(n)
            h = 0.35
            ax.barh(y - h/2, pre_means, h, label="Pre", color="#6366f1", edgecolor=cc["edge"])
            ax.barh(y + h/2, post_means, h, label="Post", color="#059669", edgecolor=cc["edge"])
            ax.set_yticks(y)
            ax.set_yticklabels(labels, fontsize=9)
            ax.set_xlabel("Mittelwert (1-5)", fontsize=10)
            ax.set_xlim(1, 5)
            ax.invert_yaxis()
            ax.legend(fontsize=9, facecolor=cc["legend_bg"], edgecolor=cc["legend_border"],
                      labelcolor=cc["legend_text"])
            fig.tight_layout()

            card = W.make_card(self._viz_frame)
            card.pack(fill="x", padx=8, pady=4)
            canvas = FigureCanvasTkAgg(fig, master=card)
            canvas.draw()
            canvas.get_tk_widget().pack(fill="x", padx=4, pady=4)
        except Exception:
            pass

    def _viz_full_stats_table(self):
        """Appendix: full statistics table."""
        self._viz_likert_table()
        self._viz_comparison_table()

    def _export_docx(self):
        """Export all sections to a DOCX file."""
        from tkinter import filedialog
        path = filedialog.asksaveasfilename(
            title="Export Report",
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx"), ("All files", "*.*")],
            initialfile="BC4D_Staffel13_Evaluation.docx",
        )
        if not path:
            return

        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Title page
            staffel = self.app.app_state.staffel_name or "13"
            title = doc.add_heading(f"Evaluierungsbericht BC4D Staffel {staffel}", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            meta = doc.add_paragraph()
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta.add_run("ISD Deutschland\n").bold = True
            meta.add_run(f"Erstellt mit BC4D Intel\n")
            import datetime
            meta.add_run(f"{datetime.date.today().strftime('%d.%m.%Y')}")

            doc.add_page_break()

            for section, label in SECTION_LABELS.items():
                text = self._sections.get(section, "")
                if not text:
                    continue

                doc.add_heading(label, level=1)

                for para_text in text.split("\n"):
                    stripped = para_text.strip()
                    if not stripped:
                        continue

                    # Detect markdown-style headings
                    if stripped.startswith("### "):
                        doc.add_heading(stripped[4:], level=3)
                    elif stripped.startswith("## "):
                        doc.add_heading(stripped[3:], level=2)
                    elif stripped.startswith("- ") or stripped.startswith("* "):
                        p = doc.add_paragraph(stripped[2:], style="List Bullet")
                    elif stripped.startswith("1. ") or stripped.startswith("2. ") or stripped.startswith("3. "):
                        p = doc.add_paragraph(stripped[3:], style="List Number")
                    else:
                        # Handle bold markers
                        p = doc.add_paragraph()
                        parts = stripped.split("**")
                        for i, part in enumerate(parts):
                            if part:
                                run = p.add_run(part)
                                if i % 2 == 1:  # odd parts are bold
                                    run.bold = True

            # Add data tables to relevant sections
            self._docx_add_tables(doc)

            doc.save(path)
            self._status_lbl.configure(
                text=f"Exported to {path.split('/')[-1]}", text_color=C.SUCCESS)
        except ImportError:
            self._status_lbl.configure(
                text="python-docx not installed. Run: pip install python-docx",
                text_color=C.DANGER)
        except Exception as e:
            self._status_lbl.configure(text=f"Export error: {e}", text_color=C.DANGER)

    def _docx_add_tables(self, doc):
        """Add formatted data tables to the DOCX document."""
        from docx.shared import Pt, Inches, RGBColor
        from docx.oxml.ns import qn

        def add_table(doc, title, headers, rows):
            doc.add_heading(title, level=3)
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Light Grid Accent 1"
            # Header
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = h
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(9)
            # Data
            for i, row_data in enumerate(rows):
                for j, val in enumerate(row_data):
                    cell = table.rows[i + 1].cells[j]
                    cell.text = str(val)
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(9)
            doc.add_paragraph("")  # spacer

        pre_df, pre_roles, post_df, post_roles = self._get_import_data()

        # Quantitative stats table
        if post_df is not None and post_roles:
            try:
                from bc4d_intel.core.stats_engine import analyze_all_likert
                items = analyze_all_likert(post_df, post_roles)
                rows = []
                for item in items:
                    s = item["stats"]
                    if s["mean"] is not None:
                        rows.append([item["label"][:50], str(s["n"]),
                                     str(s["mean"]), str(s["sd"]),
                                     f"{s['pct_agree']}%"])
                if rows:
                    add_table(doc, "Nachbefragung: Likert-Skalen",
                              ["Item", "N", "M", "SD", "Zustimmung"], rows)
            except Exception:
                pass

        # Matched comparison table
        match_result = getattr(self.app, "_match_result", None)
        if match_result and pre_roles and post_roles:
            matched = match_result.get("matched")
            if matched is not None and len(matched) > 0:
                try:
                    from bc4d_intel.core.stats_engine import analyze_matched_likert
                    comparisons = analyze_matched_likert(matched, pre_roles, post_roles)
                    rows = []
                    for c in comparisons:
                        comp = c.get("comparison", {})
                        if "error" in comp:
                            continue
                        sig = "*" if comp.get("significant") else ""
                        rows.append([
                            c["label"][:40],
                            str(comp["pre_mean"]), str(comp["post_mean"]),
                            f"{'+' if comp['mean_change'] > 0 else ''}{comp['mean_change']}",
                            str(comp["cohens_d"]),
                            f"{comp['effect_label']}{sig}",
                        ])
                    if rows:
                        add_table(doc, "Pre/Post-Vergleich (gematchtes Panel)",
                                  ["Item", "Pre M", "Post M", "Diff", "d", "Effekt"],
                                  rows)
                except Exception:
                    pass

        # Qualitative category tables
        results = getattr(self.app, "_analysis_results", {})
        if not results:
            tagged = self.app.app_state.tagged_responses
            if tagged:
                results = {k: {"classifications": v} for k, v in tagged.items()}

        if results:
            doc.add_heading("Qualitative Kategorien", level=2)
            from collections import Counter
            for label, data in list(results.items())[:8]:
                classifications = data.get("classifications", [])
                if not classifications:
                    continue
                total = len(classifications)
                counts = Counter(c.get("cluster_title", "?") for c in classifications)
                rows = [[cat[:40], str(n), f"{n/total*100:.0f}%"]
                        for cat, n in counts.most_common(10)]
                short = label.split("]")[-1].strip()[:50] if "]" in label else label[:50]
                add_table(doc, short, ["Kategorie", "N", "%"], rows)

    def refresh(self):
        """Sync sections with app_state. Detect if data changed since last generation."""
        # Always sync from app_state (may have been updated by re-analysis)
        if self.app.app_state.report_sections:
            self._sections = dict(self.app.app_state.report_sections)

        # Check if analysis data is newer than report sections
        n_classified = sum(len(v) for v in self.app.app_state.tagged_responses.values())
        if self._sections and n_classified > 0:
            # Show hint that data may have changed
            self._status_lbl.configure(
                text=f"Data: {n_classified} classified responses. "
                     f"Re-generate sections to include latest changes.",
                text_color=C.MUTED)

    def rebuild(self):
        for w in self.winfo_children():
            w.destroy()
        self.configure(fg_color=C.BG)
        self._build()
