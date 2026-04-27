import uvicorn
import os, shutil, json, time, asyncio, random
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.responses import RedirectResponse, JSONResponse, StreamingResponse, FileResponse, Response
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, List

from bc4d_intel.app_state import AppState
from bc4d_intel.core.data_loader import load_survey
from bc4d_intel.core.panel_matcher import match_panels
from bc4d_intel.constants import SESSION_DIR


class ReassignRequest(BaseModel):
    question: str
    response_index: int
    new_cluster_id: str
    new_cluster_title: str
    new_main_category: str


class AddCategoryRequest(BaseModel):
    question: str
    main_category: str
    sub_category: str


class RenameClusterRequest(BaseModel):
    question: str
    cluster_id: str
    new_title: str


class GenerateReportRequest(BaseModel):
    sections: List[str]


class ExportPayload(BaseModel):
    sections: Optional[dict] = None


app = FastAPI(title="BC4D Intel Web App")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], allow_credentials=True,
    allow_methods=["*"], allow_headers=["*"],
)

static_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static")
os.makedirs(static_dir, exist_ok=True)
app.mount("/static", StaticFiles(directory=static_dir), name="static")

# ── In-memory match result (cleared on session clear) ─────────────────────────
_match_result: Optional[dict] = None

_PARQUET_PRE     = os.path.join(SESSION_DIR, "df_pre.parquet")
_PARQUET_POST    = os.path.join(SESSION_DIR, "df_post.parquet")
_PARQUET_MATCHED = os.path.join(SESSION_DIR, "df_matched.parquet")


def _save_match_result(result: dict):
    """Persist match result DataFrames so they survive uvicorn hot-reloads."""
    os.makedirs(SESSION_DIR, exist_ok=True)
    try:
        if result.get("pre_all") is not None:   result["pre_all"].to_parquet(_PARQUET_PRE)
        if result.get("post_all") is not None:  result["post_all"].to_parquet(_PARQUET_POST)
        if result.get("matched") is not None:   result["matched"].to_parquet(_PARQUET_MATCHED)
    except Exception as e:
        print(f"[warn] Could not save parquet: {e}")


def _get_match_result():
    global _match_result
    if _match_result is not None:
        return _match_result
    # Try loading parquet cache first (survives hot-reloads)
    try:
        import pandas as pd
        if os.path.exists(_PARQUET_PRE) and os.path.exists(_PARQUET_POST):
            state = AppState.load()
            pre_df     = pd.read_parquet(_PARQUET_PRE)
            post_df    = pd.read_parquet(_PARQUET_POST)
            matched_df = pd.read_parquet(_PARQUET_MATCHED) if os.path.exists(_PARQUET_MATCHED) else None
            _match_result = {
                "pre_all":  pre_df,
                "post_all": post_df,
                "matched":  matched_df,
                "stats": {
                    "n_pre_total":     state.n_pre,
                    "n_post_total":    state.n_post,
                    "n_matched":       state.matched_pairs,
                    "n_pre_only":      state.n_pre - state.matched_pairs,
                    "n_post_only":     state.n_post - state.matched_pairs,
                    "match_rate_pre":  round(state.matched_pairs / max(state.n_pre, 1) * 100, 1),
                    "match_rate_post": round(state.matched_pairs / max(state.n_post, 1) * 100, 1),
                    "error": "",
                },
            }
            return _match_result
    except Exception as e:
        print(f"[warn] Parquet load failed: {e}")
    # Fall back to re-running match from Excel files
    state = AppState.load()
    if state.pre_survey_path and state.post_survey_path and \
       os.path.exists(state.pre_survey_path) and os.path.exists(state.post_survey_path):
        try:
            pre_df,  pre_roles  = load_survey(state.pre_survey_path)
            post_df, post_roles = load_survey(state.post_survey_path)
            state.pre_columns  = pre_roles
            state.post_columns = post_roles
            _match_result = match_panels(pre_df, pre_roles, post_df, post_roles)
            _save_match_result(_match_result)
        except Exception as e:
            print(f"[warn] Could not reload match result: {e}")
    return _match_result



# ══════════════════════════════════════════════════════════════════════════════
# ROOT
# ══════════════════════════════════════════════════════════════════════════════
@app.get("/")
def redirect_to_app():
    return RedirectResponse(url="/static/index.html")


# ══════════════════════════════════════════════════════════════════════════════
# STATE
# ══════════════════════════════════════════════════════════════════════════════
@app.get("/api/state")
def get_state():
    state  = AppState.load()
    result = _get_match_result()
    n_q   = len(state.tagged_responses)
    n_r   = sum(len(v) for v in state.tagged_responses.values())
    n_cats = sum(len(ft) for ft in state.flat_taxonomies.values())
    n_high = sum(1 for vs in state.tagged_responses.values()
                 for c in vs if c.get("confidence") == "high")
    pct_high = round(n_high / max(n_r, 1) * 100) if n_r else 0

    stats = result["stats"] if result else {}
    return {
        "n_pre":            state.n_pre,
        "n_post":           state.n_post,
        "matched_pairs":    state.matched_pairs,
        "n_matched":        stats.get("n_matched", state.matched_pairs),
        "n_unmatched_pre":  stats.get("n_pre_only", 0),
        "n_unmatched_post": stats.get("n_post_only", 0),
        "match_rate_pct":   stats.get("match_rate_post", 0),
        "tagged_questions": n_q,
        "n_questions":      n_q,
        "n_categories":     n_cats,
        "n_responses":      n_r,
        "n_high_confidence": n_high,
        "pct_high":         pct_high,
        "has_match_result": result is not None,
        "has_results":      n_q > 0,
        "staffel_name":     state.staffel_name,
        "has_api_key":      bool(state.api_key),
        "warnings":         [],
    }


# ══════════════════════════════════════════════════════════════════════════════
# DEMO / SIMULATION MODE
# ══════════════════════════════════════════════════════════════════════════════
@app.post("/api/demo/load")
def load_demo_data():
    """Populate session with realistic synthetic survey data so all charts render."""
    global _match_result
    state = AppState.load()

    # ── Survey counts ──
    state.n_pre         = 304
    state.n_post        = 152
    state.matched_pairs = 130
    state.staffel_name  = "Staffel 13 (Demo)"

    # ── Synthetic match result DataFrames ──
    import pandas as pd, numpy as np
    rng = np.random.default_rng(42)

    pre_questions = [
        "Ich traue mir zu, Falschinformationen im Netz zu erkennen.",
        "Ich kenne Strategien, um Hassrede online zu begegnen.",
        "Ich fühle mich sicher im Umgang mit digitalen Medien.",
        "Ich diskutiere Medienthemen in meinem sozialen Umfeld.",
        "Das Thema Desinformation ist für meinen Alltag relevant.",
        "Ich überprüfe Quellen, bevor ich Inhalte teile.",
        "Ich fühle mich kompetent, andere über Desinformation aufzuklären.",
    ]
    # Use SAME question labels for pre & post so fuzzy matching finds pairs
    # The matched DF will have col_pre and col_post columns
    post_questions = pre_questions   # identical labels so ratio=1.0 in fuzzy match

    def likert(n, mu, sd): return np.clip(rng.normal(mu, sd, n).round().astype(int), 1, 5)

    n_pre = 304; n_post = 152; n_match = 130

    pre_data = {q: likert(n_pre, mu, 1.1) for q, mu in zip(pre_questions, [2.8,2.5,3.1,3.3,3.7,3.0,2.6])}
    pre_df = pd.DataFrame(pre_data)

    post_data = {q: likert(n_post, mu, 1.0) for q, mu in zip(post_questions, [3.8,3.6,4.0,3.9,4.1,3.9,3.7])}
    post_df = pd.DataFrame(post_data)

    # matched: columns named <question>_pre and <question>_post (what analyze_matched_likert expects)
    matched_rows = {
        **{q + "_pre":  likert(n_match, mu, 1.1) for q, mu in zip(pre_questions,  [2.8,2.5,3.1,3.3,3.7,3.0,2.6])},
        **{q + "_post": likert(n_match, mu, 1.0) for q, mu in zip(post_questions, [3.8,3.6,4.0,3.9,4.1,3.9,3.7])},
    }
    matched_df = pd.DataFrame(matched_rows)

    pre_roles  = {q: "likert" for q in pre_questions}
    post_roles = {q: "likert" for q in post_questions}
    state.pre_columns  = pre_roles
    state.post_columns = post_roles

    _match_result = {
        "pre_all":  pre_df,
        "post_all": post_df,
        "matched":  matched_df,
        "stats": {
            "n_pre_total": n_pre, "n_post_total": n_post,
            "n_matched": n_match, "n_pre_only": n_pre - n_match,
            "n_post_only": n_post - n_match,
            "match_rate_pre": round(n_match / n_pre * 100, 1),
            "match_rate_post": round(n_match / n_post * 100, 1),
            "error": "",
        },
    }
    _save_match_result(_match_result)  # persist to parquet — survives uvicorn hot-reloads

    # ── Synthetic AI analysis results ──
    DEMO_QUESTIONS = [
        "[Pre] Warum denken Sie so?",
        "[Post] Was hat sich für Sie durch das Training verändert?",
        "[Pre] Wie nutzen Sie soziale Medien?",
        "[Post] Was nehmen Sie aus dem Kurs mit?",
    ]
    CATS = [
        {"main": "Medienkompetenz", "subs": ["Quellenkritik", "Faktenchecks", "Algorithmen-Verständnis"]},
        {"main": "Persönliche Haltung", "subs": ["Selbstwirksamkeit", "Skepsis", "Verantwortungsbewusstsein"]},
        {"main": "Soziale Dimension", "subs": ["Weitergabe von Wissen", "Diskussion im Alltag", "Netzwerk-Effekte"]},
    ]
    SAMPLE_TEXTS = [
        "Ich prüfe jetzt viel häufiger die Quellen bevor ich etwas teile.",
        "Das Training hat mir gezeigt wie leicht man manipuliert werden kann.",
        "Ich bin jetzt kritischer gegenüber Nachrichten in sozialen Medien.",
        "Es war überraschend, wie viele Faktoren Algorithmen beeinflussen.",
        "Das Thema war mir vor dem Training kaum bekannt.",
        "Ich spreche jetzt öfter mit Familie über Desinformation.",
        "Die Übungen haben mir sehr geholfen, Beispiele zu erkennen.",
        "Bis jetzt habe ich kaum darüber nachgedacht wie Medien funktionieren.",
        "Sehr aufschlussreich, ich empfehle es jedem.",
        "Die Gruppenarbeiten haben meinen Blickwinkel erweitert.",
        "Ich war skeptisch, aber das Training war wirklich gut.",
        "Besonders beeindruckt haben mich die Fallstudien.",
        "Man merkt jetzt viel schneller wenn etwas manipulativ ist.",
        "Ich teile jetzt deutlich weniger ungeprüfte Inhalte.",
        "Das Training sollte in Schulen Pflicht sein.",
    ]

    state.tagged_responses = {}
    state.taxonomies       = {}
    state.flat_taxonomies  = {}

    for qi, q_label in enumerate(DEMO_QUESTIONS):
        flat_tax = []
        all_subs = []
        for main_obj in CATS:
            for sub in main_obj["subs"]:
                sub_id = f"cat_{len(flat_tax)}"
                flat_tax.append({"id": sub_id, "title": sub, "main_category": main_obj["main"], "count": 0})
                all_subs.append((sub_id, sub, main_obj["main"]))

        taxonomy = {"categories": [
            {"main_category": mc["main"], "sub_categories": [
                {"id": ft["id"], "title": ft["title"],
                 "include_rule": f"Antworten, die {ft['title'].lower()} thematisieren.",
                 "exclude_rule": "",
                 "examples": [SAMPLE_TEXTS[i % len(SAMPLE_TEXTS)] for i in range(2)]}
                for ft in flat_tax if ft["main_category"] == mc["main"]
            ]}
            for mc in CATS
        ]}

        n_resp = rng.integers(20, 50)
        classifications = []
        for i in range(n_resp):
            sub_id, sub_title, main_cat = all_subs[i % len(all_subs)]
            text = SAMPLE_TEXTS[(qi * 5 + i) % len(SAMPLE_TEXTS)]
            conf = rng.choice(["high", "high", "medium", "low"])
            classifications.append({
                "text": text, "cluster_id": sub_id,
                "cluster_title": sub_title, "main_category": main_cat,
                "confidence": conf, "human_override": "",
            })
            flat_tax_item = next(f for f in flat_tax if f["id"] == sub_id)
            flat_tax_item["count"] += 1

        state.tagged_responses[q_label] = classifications
        state.taxonomies[q_label]       = taxonomy
        state.flat_taxonomies[q_label]  = flat_tax

    state.save()
    n_r = sum(len(v) for v in state.tagged_responses.values())
    return {
        "status": "ok",
        "message": "Demo data loaded. All charts are now available.",
        "n_pre": state.n_pre, "n_post": state.n_post,
        "matched_pairs": state.matched_pairs,
        "n_questions": len(state.tagged_responses),
        "n_responses": n_r,
    }


# ══════════════════════════════════════════════════════════════════════════════
# UPLOAD & MATCHING
# ══════════════════════════════════════════════════════════════════════════════
@app.post("/api/upload")
async def upload_surveys(
    pre_file:  UploadFile = File(...),
    post_file: UploadFile = File(...),
):
    global _match_result
    state = AppState.load()
    os.makedirs(SESSION_DIR, exist_ok=True)

    pre_path  = os.path.join(SESSION_DIR, "temp_pre.xlsx")
    post_path = os.path.join(SESSION_DIR, "temp_post.xlsx")

    with open(pre_path,  "wb") as f: shutil.copyfileobj(pre_file.file,  f)
    with open(post_path, "wb") as f: shutil.copyfileobj(post_file.file, f)

    try:
        pre_df,  pre_roles  = load_survey(pre_path)
        post_df, post_roles = load_survey(post_path)
        result = match_panels(pre_df, pre_roles, post_df, post_roles)
        _match_result = result
        _save_match_result(result)

        stats = result["stats"]
        state.pre_survey_path  = pre_path
        state.post_survey_path = post_path
        state.n_pre            = stats.get("n_pre_total",  len(pre_df))
        state.n_post           = stats.get("n_post_total", len(post_df))
        state.matched_pairs    = stats.get("n_matched", 0)
        state.pre_columns      = pre_roles
        state.post_columns     = post_roles
        state.save()

        return {
            "status":           "success",
            "n_pre":            state.n_pre,
            "n_post":           state.n_post,
            "n_matched":        stats.get("n_matched", 0),
            "n_unmatched_pre":  stats.get("n_pre_only", 0),
            "n_unmatched_post": stats.get("n_post_only", 0),
            "match_rate_pct":   stats.get("match_rate_post", 0),
            "warnings":         [],
        }
    except Exception as e:
        import traceback; traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# ══════════════════════════════════════════════════════════════════════════════
# DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════
@app.get("/api/dashboard")
def get_dashboard():
    state  = AppState.load()
    result = _get_match_result()
    stats  = result["stats"] if result else {}
    return {
        "n_pre":             state.n_pre,
        "n_post":            state.n_post,
        "n_matched":         stats.get("n_matched",      state.matched_pairs),
        "n_unmatched_pre":   stats.get("n_pre_only",     0),
        "n_unmatched_post":  stats.get("n_post_only",    0),
        "match_rate_pct":    stats.get("match_rate_post", 0),
    }


def _build_likert_items(df, roles):
    """Return Likert item stats from a DataFrame."""
    if df is None or not roles:
        return []
    from bc4d_intel.core.stats_engine import analyze_all_likert
    items = analyze_all_likert(df, roles)
    out = []
    for item in items:
        s = item["stats"]
        if s.get("mean") is not None:
            out.append({"label": item["label"], "stats": {
                "n":           s["n"],
                "mean":        s["mean"],
                "sd":          s["sd"],
                "pct_agree":   s.get("pct_agree", 0),
                "pct_disagree":s.get("pct_disagree", 0),
                "distribution":s.get("distribution", {}),
            }})
    return out


def _build_demographic_items(df, roles):
    if df is None or not roles:
        return []
    from bc4d_intel.core.stats_engine import analyze_demographics
    return analyze_demographics(df, roles)


@app.get("/api/dashboard/pre")
def dashboard_pre():
    result = _get_match_result()
    state  = AppState.load()
    if not result:
        return {"items": [], "demographics": [], "n": 0}
    pre_df = result.get("pre_all")
    items  = _build_likert_items(pre_df, state.pre_columns)
    demographics = _build_demographic_items(pre_df, state.pre_columns)
    return {"items": items, "demographics": demographics, "n": len(pre_df) if pre_df is not None else 0}


@app.get("/api/dashboard/post")
def dashboard_post():
    result = _get_match_result()
    state  = AppState.load()
    if not result:
        return {"items": [], "demographics": [], "n": 0}
    post_df = result.get("post_all")
    items   = _build_likert_items(post_df, state.post_columns)
    demographics = _build_demographic_items(post_df, state.post_columns)
    return {"items": items, "demographics": demographics, "n": len(post_df) if post_df is not None else 0}


@app.get("/api/dashboard/matched")
def dashboard_matched():
    result = _get_match_result()
    state  = AppState.load()
    if not result:
        return {"comparisons": []}
    matched = result.get("matched")
    if matched is None or len(matched) == 0:
        return {"comparisons": []}
    from bc4d_intel.core.stats_engine import analyze_matched_likert
    comparisons = analyze_matched_likert(matched, state.pre_columns, state.post_columns)
    out = []
    for c in (comparisons or []):
        comp = c.get("comparison", {})
        if "error" in comp:
            continue
        out.append({
            "label": c["label"],
            "comparison": {
                "pre_mean":    comp.get("pre_mean"),
                "post_mean":   comp.get("post_mean"),
                "mean_change": comp.get("mean_change"),
                "cohens_d":    comp.get("cohens_d"),
                "effect_label":comp.get("effect_label", ""),
                "significant": comp.get("significant", False),
                "p_value":     comp.get("p_value"),
                "direction":   comp.get("direction", ""),
                "improved_pct":comp.get("improved_pct"),
            },
        })
    return {"comparisons": out}


# ══════════════════════════════════════════════════════════════════════════════
# RAW DATA EXPORT (NEW)
# ══════════════════════════════════════════════════════════════════════════════
@app.get("/api/export/data/{export_type}")
def export_raw_data(export_type: str, question: Optional[str] = None):
    result = _get_match_result()
    if not result and export_type != "clusters":
        raise HTTPException(status_code=400, detail="Data not loaded yet.")
        
    df = None
    filename = "bc4d_data.csv"
    
    if export_type == "pre":
        df = result.get("pre_all")
        filename = "bc4d_pre_survey.csv"
    elif export_type == "post":
        df = result.get("post_all")
        filename = "bc4d_post_survey.csv"
    elif export_type == "matched":
        df = result.get("matched")
        filename = "bc4d_matched_panel.csv"
    elif export_type == "clusters":
        state = AppState.load()
        if not question:
            raise HTTPException(status_code=400, detail="Question parameter required for clusters.")
        classifications = state.tagged_responses.get(question, [])
        if not classifications:
            raise HTTPException(status_code=404, detail="No clusters found for this question.")
        import pandas as pd
        df = pd.DataFrame(classifications)
        filename = "bc4d_clusters_data.csv"
    else:
        raise HTTPException(status_code=400, detail="Invalid export type.")
        
    if df is None or len(df) == 0:
        raise HTTPException(status_code=404, detail="No data available.")
        
    csv_str = df.to_csv(index=False)
    headers = {
        "Content-Disposition": f"attachment; filename={filename}"
    }
    return Response(content=csv_str, media_type="text/csv", headers=headers)


# ══════════════════════════════════════════════════════════════════════════════
# DATA MUTATION & OVERRIDES (NEW)
# ══════════════════════════════════════════════════════════════════════════════
def _recalc_flat_taxonomy_counts(state: AppState, question: str):
    if question not in state.tagged_responses or question not in state.flat_taxonomies:
        return
    counts = {}
    for r in state.tagged_responses[question]:
        cid = r.get("human_override") or r.get("cluster_id")
        counts[cid] = counts.get(cid, 0) + 1
    for cat in state.flat_taxonomies[question]:
        cat["count"] = counts.get(cat["id"], 0)


@app.post("/api/responses/reassign")
def api_reassign_response(req: ReassignRequest):
    state = AppState.load()
    if req.question not in state.tagged_responses:
        raise HTTPException(status_code=404, detail="Question not found")
        
    responses = state.tagged_responses[req.question]
    if req.response_index < 0 or req.response_index >= len(responses):
        raise HTTPException(status_code=400, detail="Invalid response index")
        
    r = responses[req.response_index]
    r["human_override"] = req.new_cluster_id
    r["main_category"] = req.new_main_category
    r["cluster_title"] = req.new_cluster_title
    
    _recalc_flat_taxonomy_counts(state, req.question)
    state.save()
    return {"status": "ok"}


@app.post("/api/responses/category/add")
def api_add_category(req: AddCategoryRequest):
    state = AppState.load()
    if req.question not in state.taxonomies:
        raise HTTPException(status_code=404, detail="Question not found")
        
    flat_tax = state.flat_taxonomies.get(req.question, [])
    existing_ids = {c["id"] for c in flat_tax}
    new_id = f"cat_custom_{len(existing_ids) + 1}"
    
    # Add to flat
    flat_tax.append({
        "id": new_id,
        "title": req.sub_category,
        "main_category": req.main_category,
        "description": "User added",
        "count": 0
    })
    state.flat_taxonomies[req.question] = flat_tax
    
    # Add to hierarchical
    taxonomy = state.taxonomies[req.question]
    found_main = False
    if "categories" not in taxonomy:
        taxonomy["categories"] = []
    
    for mc in taxonomy["categories"]:
        if mc.get("main_category") == req.main_category:
            mc.setdefault("sub_categories", []).append({
                "id": new_id, "title": req.sub_category, "examples": [],
                "include_rule": "User added", "exclude_rule": ""
            })
            found_main = True
            break
            
    if not found_main:
        taxonomy["categories"].append({
            "id": f"main_{new_id}",
            "main_category": req.main_category,
            "sub_categories": [{
                "id": new_id, "title": req.sub_category, "examples": [],
                "include_rule": "User added", "exclude_rule": ""
            }]
        })
        
    state.save()
    return {"status": "ok", "new_id": new_id}


@app.post("/api/clusters/rename")
def api_rename_cluster(req: RenameClusterRequest):
    state = AppState.load()
    if req.question not in state.taxonomies:
        raise HTTPException(status_code=404, detail="Question not found")
        
    # Update hierarchical
    for mc in state.taxonomies[req.question].get("categories", []):
        for sub in mc.get("sub_categories", []):
            if sub["id"] == req.cluster_id:
                sub["title"] = req.new_title
                
    # Update flat
    for flat in state.flat_taxonomies.get(req.question, []):
        if flat["id"] == req.cluster_id:
            flat["title"] = req.new_title
            
    # Update responses
    for r in state.tagged_responses.get(req.question, []):
        if (r.get("human_override") or r.get("cluster_id")) == req.cluster_id:
            r["cluster_title"] = req.new_title
            
    state.save()
    return {"status": "ok"}


# ══════════════════════════════════════════════════════════════════════════════
# AI ENGINE
# ══════════════════════════════════════════════════════════════════════════════
@app.get("/api/analysis/estimate")
def analysis_estimate():
    state  = AppState.load()
    result = _get_match_result()
    if not result:
        return {"text": "Load survey data first (or load demo data) to see estimates.", "has_results": False}

    n_questions = 0; total_responses = 0
    for survey_type, df_key, roles in [
        ("Pre",  "pre_all",  state.pre_columns),
        ("Post", "post_all", state.post_columns),
    ]:
        df = result.get(df_key)
        if df is None: continue
        for col, role in roles.items():
            if role == "free_text" and col in df.columns:
                n_resp = len([r for r in df[col].dropna().astype(str).tolist() if len(r.strip()) > 5])
                if n_resp >= 5:
                    n_questions += 1; total_responses += n_resp

    taxonomy_cost    = n_questions * 0.03
    classify_batches = (total_responses + 19) // 20
    classify_cost    = classify_batches * 0.002
    total_cost       = taxonomy_cost + classify_cost
    est_time         = n_questions * 2 + classify_batches * 2

    lines = [
        f"Found: {n_questions} open-ended questions, {total_responses} responses total",
        "",
        f"Category design: {n_questions} questions × ~$0.03 = ~${taxonomy_cost:.2f}",
        f"Classification: {total_responses} responses in {classify_batches} batches × ~$0.002 = ~${classify_cost:.3f}",
        "",
        f"Estimated total cost: ~${total_cost:.2f}",
        f"Estimated time: ~{max(est_time, 10)}s",
    ]
    state_snap = get_state()
    return {
        "text":              "\n".join(lines),
        "has_results":       state_snap["has_results"],
        "n_questions":       state_snap["n_questions"],
        "n_categories":      state_snap["n_categories"],
        "n_responses":       state_snap["n_responses"],
        "n_high_confidence": state_snap["n_high_confidence"],
        "pct_high":          state_snap["pct_high"],
    }


@app.post("/api/analysis/run")
async def run_analysis():
    """Streaming SSE endpoint for real AI analysis."""
    state  = AppState.load()
    result = _get_match_result()

    async def generate():
        if not state.api_key:
            yield f"data: {json.dumps({'error': 'Set API key in Settings first.'})}\n\n"
            return
        if not result:
            yield f"data: {json.dumps({'error': 'Load survey data first (Import screen).'})}\n\n"
            return

        try:
            from bc4d_intel.core.embedder import full_pipeline
        except ImportError:
            yield f"data: {json.dumps({'error': 'AI modules not available.'})}\n\n"
            return

        try:
            from bc4d_intel.core.answer_cache import (
                deduplicate, add_to_cache,
                get_cached_taxonomy, save_taxonomy,
                classify_with_llm,
            )
        except ImportError:
            yield f"data: {json.dumps({'error': 'Answer cache module not available.'})}\n\n"
            return

        # Build list of free-text questions
        all_ft = []
        for survey_type, df_key, roles in [
            ("Pre",  "pre_all",  state.pre_columns),
            ("Post", "post_all", state.post_columns),
        ]:
            df = result.get(df_key)
            if df is None: continue
            for col, role in roles.items():
                if role == "free_text" and col in df.columns:
                    responses = [r for r in df[col].dropna().astype(str).tolist() if len(r.strip()) > 5]
                    if len(responses) >= 5:
                        label = f"[{survey_type}] {col}"
                        all_ft.append((label, col, responses))

        total_q = len(all_ft)
        if total_q == 0:
            yield f"data: {json.dumps({'error': 'No open-ended questions detected.'})}\n\n"
            return

        all_results = {}
        for qi, (label, col_name, responses) in enumerate(all_ft):
            pct = (qi + 0.5) / max(total_q, 1)
            yield f"data: {json.dumps({'progress': pct, 'detail': f'Analyzing: {label[:40]}... ({len(responses)} responses)'})}\n\n"
            await asyncio.sleep(0)

            try:
                taxonomy = get_cached_taxonomy(label)
                if not taxonomy:
                    res = full_pipeline(responses, state.api_key, question=col_name)
                    if res.get("taxonomy"):
                        save_taxonomy(label, res["taxonomy"], n_responses=len(responses))
                        taxonomy = res["taxonomy"]
                    classified = res.get("classifications", [])
                else:
                    deduped, remaining = deduplicate(label, responses)
                    llm_classified = classify_with_llm(col_name, remaining, taxonomy, state.api_key) if remaining else []
                    classified = deduped + llm_classified

                flat_taxonomy = []
                if taxonomy:
                    for mc in taxonomy.get("categories", []):
                        for sub in mc.get("sub_categories", []):
                            count = sum(1 for c in classified if (c.get("human_override") or c.get("cluster_id","")) == sub["id"])
                            flat_taxonomy.append({
                                "id": sub["id"], "title": sub["title"],
                                "main_category": mc["main_category"],
                                "description": sub.get("include_rule",""),
                                "count": count,
                            })

                add_to_cache(label, classified, staffel=state.staffel_name)
                state.tagged_responses[label]  = classified
                state.taxonomies[label]         = taxonomy or {}
                state.flat_taxonomies[label]    = flat_taxonomy
                state.save()
                all_results[label] = {"taxonomy": taxonomy, "flat_taxonomy": flat_taxonomy, "classifications": classified}

            except Exception as e:
                yield f"data: {json.dumps({'detail': f'Error on {label[:30]}: {str(e)}'})}\n\n"
            await asyncio.sleep(0)

        n_r    = sum(len(r.get("classifications",[])) for r in all_results.values())
        n_cats = sum(len(r.get("flat_taxonomy",[]))    for r in all_results.values())
        n_high = sum(1 for r in all_results.values() for c in r.get("classifications",[]) if c.get("confidence")=="high")
        pct_h  = round(n_high / max(n_r,1)*100) if n_r else 0

        yield f"data: {json.dumps({'progress': 1.0, 'done': True, 'detail': 'Complete!', 'n_questions': len(all_results), 'n_categories': n_cats, 'n_responses': n_r, 'n_high_confidence': n_high, 'pct_high': pct_h})}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")


@app.post("/api/analysis/import")
async def import_analysis_results(file: UploadFile = File(...)):
    state = AppState.load()
    try:
        content = await file.read()
        data = json.loads(content)
        if "questions" not in data:
            raise HTTPException(status_code=400, detail="Invalid file: no 'questions' key found.")
        for label, qdata in data["questions"].items():
            state.tagged_responses[label] = qdata.get("classifications", [])
            state.taxonomies[label]       = qdata.get("taxonomy", {})
            state.flat_taxonomies[label]  = qdata.get("flat_taxonomy", [])
        state.save()
        state_snap = get_state()
        return {"status": "success", **state_snap}
    except HTTPException: raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ══════════════════════════════════════════════════════════════════════════════
# CLUSTERS  (drives both Clusters and Responses pages)
# ══════════════════════════════════════════════════════════════════════════════
@app.get("/api/clusters")
def get_clusters():
    state = AppState.load()
    return {
        label: {
            "taxonomy":        state.taxonomies.get(label, {}),
            "flat_taxonomy":   state.flat_taxonomies.get(label, []),
            "classifications": state.tagged_responses[label],
        }
        for label in state.tagged_responses
    }


# ══════════════════════════════════════════════════════════════════════════════
# SETTINGS  (GET + POST)
# ══════════════════════════════════════════════════════════════════════════════
@app.get("/api/settings")
def get_settings():
    state = AppState.load()
    return {
        "api_key":          state.api_key or "",
        "staffel_name":     state.staffel_name or "",
        "n_pre":            state.n_pre,
        "n_post":           state.n_post,
        "matched_pairs":    state.matched_pairs,
        "tagged_questions": len(state.tagged_responses),
        "report_sections":  len(state.report_sections),
    }


class SettingsPayload(BaseModel):
    api_key:      Optional[str] = None
    staffel_name: Optional[str] = None


@app.post("/api/settings")
async def save_settings(payload: SettingsPayload):
    state = AppState.load()
    message = "Saved."
    if payload.api_key is not None:
        state.api_key = payload.api_key
        try:
            from bc4d_intel.ai.claude_client import call_claude
            call_claude(system="Say OK.", user_msg="Test.", task="tagging",
                        api_key=payload.api_key, max_tokens=5)
            message = "API key saved. Connection successful ✓"
        except Exception as e:
            message = f"Saved (connection test: {str(e)[:60]})"
    if payload.staffel_name is not None:
        state.staffel_name = payload.staffel_name
    state.save()
    return {"status": "success", "message": message}


@app.post("/api/session/clear")
def clear_session():
    global _match_result
    state   = AppState.load()
    api_key = state.api_key
    from bc4d_intel.app_state import AppState as AS
    new_state = AS(api_key=api_key)
    new_state.save()
    _match_result = None
    for f in [_PARQUET_PRE, _PARQUET_POST, _PARQUET_MATCHED]:
        try: os.remove(f)
        except OSError: pass
    return {"status": "success"}


# ══════════════════════════════════════════════════════════════════════════════
# REPORT
# ══════════════════════════════════════════════════════════════════════════════
@app.get("/api/report/sections")
def get_report_sections():
    return AppState.load().report_sections or {}


@app.post("/api/report/generate/{section_key}")
async def generate_section_endpoint(section_key: str):
    """Streaming text generation for a single report section."""
    state = AppState.load()
    if not state.api_key:
        raise HTTPException(status_code=400, detail="Set API key in Settings first.")

    async def stream():
        try:
            from bc4d_intel.ai.report_writer import generate_section, build_data_context
            result  = _get_match_result()
            pre_df  = result.get("pre_all")  if result else None
            post_df = result.get("post_all") if result else None

            tagged = None
            if state.tagged_responses:
                tagged = {}
                for label, classifications in state.tagged_responses.items():
                    tagged[label] = [
                        {
                            "text":           c.get("text",""),
                            "tag":            c.get("cluster_title",""),
                            "human_override": c.get("human_override",""),
                        }
                        for c in classifications
                    ]

            context = build_data_context(
                state, result, tagged,
                pre_roles  = state.pre_columns,
                post_roles = state.post_columns,
                pre_df=pre_df, post_df=post_df,
            )

            chunks = []
            def on_stream(chunk): chunks.append(chunk)

            text = generate_section(section_key, context, state.api_key, stream_cb=on_stream)
            state.report_sections[section_key] = text
            state.save()

            for c in (chunks if chunks else [text]):
                yield c
                await asyncio.sleep(0)

        except Exception as e:
            yield f"\n\n[Error generating section: {e}]"

    return StreamingResponse(stream(), media_type="text/plain")




@app.post("/api/report/export")
async def export_report_docx(payload: ExportPayload):
    state = AppState.load()
    os.makedirs(SESSION_DIR, exist_ok=True)
    report_path = os.path.join(SESSION_DIR, "BC4D_Intel_Report.docx")

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import datetime

        doc = Document()
        staffel = state.staffel_name or "13"
        title = doc.add_heading(f"Evaluierungsbericht BC4D Staffel {staffel}", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run("ISD Deutschland\n").bold = True
        meta.add_run(f"Erstellt mit BC4D Intel — {datetime.date.today().strftime('%d.%m.%Y')}")
        doc.add_page_break()

        SECTION_LABELS = {
            "executive_summary":    "1. Executive Summary",
            "method_sample":        "2. Method & Sample",
            "quantitative_results": "3. Quantitative Results",
            "qualitative_findings": "4. Qualitative Findings",
            "pre_post_comparison":  "5. Pre/Post Comparison",
            "recommendations":      "6. Recommendations",
            "appendix":             "7. Appendix",
        }

        sections_to_export = payload.sections or state.report_sections or {}
        for key, label in SECTION_LABELS.items():
            text = sections_to_export.get(key, "")
            if not text: continue
            doc.add_heading(label, level=1)
            for line in text.split("\n"):
                stripped = line.strip()
                if not stripped: continue
                if stripped.startswith("### "): doc.add_heading(stripped[4:], level=3); continue
                if stripped.startswith("## "):  doc.add_heading(stripped[3:], level=2); continue
                if stripped.startswith("- ") or stripped.startswith("* "):
                    doc.add_paragraph(stripped[2:], style="List Bullet"); continue
                p = doc.add_paragraph()
                parts = stripped.split("**")
                for i, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        if i % 2 == 1: run.bold = True

        doc.save(report_path)
        return FileResponse(report_path, filename="BC4D_Evaluation_Report.docx",
                            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed: pip install python-docx")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
