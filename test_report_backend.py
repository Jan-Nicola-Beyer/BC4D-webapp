"""Backend test for report writing + chart data export.

Tests:
1. Data context building (stats, qualitative summaries)
2. Report section generation (mocked Sonnet)
3. Chart data export to formatted CSV/Excel
4. All 7 sections produce structured content
"""

from __future__ import annotations
import os, sys, json, time
from unittest.mock import patch

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "bc4d_intel"))
sys.path.insert(0, os.path.dirname(__file__))

RESULTS = {}


def record(name, passed, details=""):
    RESULTS[name] = {"passed": passed, "details": details}
    print(f"  [{'PASS' if passed else 'FAIL'}] {name}")
    if details:
        for line in details.split("\n")[:3]:
            print(f"         {line}")


# ════════════════════════════════════════════════════════════════
print("\n" + "=" * 70)
print("REPORT SYSTEM BACKEND TEST")
print("=" * 70)

# ── 1. Load data ────────────────────────────────────────────────
print("\n--- 1. Load Data ---\n")

from bc4d_intel.core.data_loader import load_survey
from bc4d_intel.core.panel_matcher import match_panels
from bc4d_intel.app_state import AppState

pre_df, pre_roles = load_survey("NEGATIV_Vorbefragung_Staffel13.xlsx")
post_df, post_roles = load_survey("NEGATIV_Abschlussbefragung_Staffel13.xlsx")
match_result = match_panels(pre_df, pre_roles, post_df, post_roles)

app_state = AppState()
app_state.staffel_name = "13 (NEGATIV Test)"
app_state.n_pre = len(pre_df)
app_state.n_post = len(post_df)
app_state.matched_pairs = len(match_result.get("matched", []))
app_state.unmatched_pre = app_state.n_pre - app_state.matched_pairs
app_state.unmatched_post = app_state.n_post - app_state.matched_pairs

record("data_loaded", True,
       f"Pre: {len(pre_df)}, Post: {len(post_df)}, "
       f"Matched: {app_state.matched_pairs}")

# ── 2. Build data context ──────────────────────────────────────
print("\n--- 2. Data Context ---\n")

from bc4d_intel.ai.report_writer import build_data_context

context = build_data_context(
    app_state, match_result=match_result,
    pre_roles=pre_roles, post_roles=post_roles,
    pre_df=match_result.get("pre_all", pre_df),
    post_df=match_result.get("post_all", post_df),
)

has_staffel = "Staffel:" in context
has_quant = "QUANTITATIVE" in context
has_matched = "PRE/POST" in context
has_mean = "M=" in context
has_cohens = "Cohen" in context

record("context_has_sections",
       has_staffel and has_quant and has_matched,
       f"Staffel: {has_staffel}, Quant: {has_quant}, "
       f"Matched: {has_matched}")

record("context_has_stats",
       has_mean and has_cohens,
       f"Means: {has_mean}, Cohen's d: {has_cohens}")

record("context_length", len(context) > 500,
       f"{len(context)} chars, {context.count(chr(10))} lines")

# Show a sample
print(f"\n  Context preview (first 500 chars):")
print(f"  {context[:500]}")

# ── 3. Report generation (mocked) ──────────────────────────────
print("\n\n--- 3. Report Section Generation (mocked) ---\n")

from bc4d_intel.ai.report_writer import generate_section
from bc4d_intel.ai.prompts import REPORT_SECTIONS

MOCK_RESPONSES = {
    "executive_summary": """## Zusammenfassung

Die Evaluation der BC4D-Staffel 13 umfasst N=304 Teilnehmende im Vorabfragebogen
und N=147 in der Nachbefragung (Match-Rate: 86%). Die zentralen Ergebnisse zeigen:

1. **Hohe Zufriedenheit** mit der Schulung (M=4.2, SD=0.8)
2. **Signifikante Kompetenzsteigerung** bei der Erkennung von Hassrede (d=0.65, p<.001)
3. **Praxistransfer**: 72% der Teilnehmenden haben Gelerntes bereits angewendet

Die Ergebnisse belegen die Wirksamkeit des BC4D-Programms.""",

    "method_sample": """## Methodik und Stichprobe

### Evaluationsdesign
Die Evaluation basiert auf einem Pre-Post-Design mit pseudonymisierter Zuordnung.

### Stichprobe
- Vorabfragebogen: N=304 Teilnehmende
- Nachbefragung: N=147 Teilnehmende
- Gematchte Paare: N=127 (Match-Rate: 86%)

### Drei Analyseebenen
1. Alle Pre-Befragten (Baseline)
2. Alle Post-Befragten (Ergebnisse)
3. Gematchtes Panel (individuelle Veraenderung)""",
}


def mock_sonnet(system, user_msg, task="report", api_key="", max_tokens=2500, stream_cb=None):
    """Return realistic German report text for each section."""
    for section_name, prompt in REPORT_SECTIONS.items():
        if prompt[:50] in user_msg:
            text = MOCK_RESPONSES.get(section_name,
                f"## {section_name}\n\nDieser Abschnitt wird generiert basierend auf "
                f"den bereitgestellten Daten. Die Staffel umfasst "
                f"{304} Teilnehmende im Pre-Survey.")
            if stream_cb:
                for word in text.split():
                    stream_cb(word + " ")
            return text
    return "## Abschnitt\n\nInhalt wird generiert."


all_sections_ok = True
for section_name in REPORT_SECTIONS:
    with patch("bc4d_intel.ai.report_writer.call_claude", side_effect=mock_sonnet):
        text = generate_section(section_name, context, "mock-key")

    ok = len(text) > 20 and "##" in text
    if not ok:
        all_sections_ok = False
    record(f"section_{section_name}",
           ok, f"{len(text)} chars: {text[:80]}...")

record("all_sections_generated", all_sections_ok, "")

# ── 4. DOCX export test ────────────────────────────────────────
print("\n--- 4. DOCX Export ---\n")

try:
    from docx import Document
    doc = Document()
    doc.add_heading("BC4D Evaluation Report - Staffel 13", level=0)
    for section_name in REPORT_SECTIONS:
        with patch("bc4d_intel.ai.report_writer.call_claude", side_effect=mock_sonnet):
            text = generate_section(section_name, context, "mock-key")
        doc.add_heading(section_name.replace("_", " ").title(), level=1)
        for para in text.split("\n"):
            para = para.strip()
            if not para:
                continue
            if para.startswith("## "):
                doc.add_heading(para[3:], level=2)
            elif para.startswith("### "):
                doc.add_heading(para[4:], level=3)
            elif para.startswith("- ") or para.startswith("* "):
                doc.add_paragraph(para[2:], style="List Bullet")
            else:
                doc.add_paragraph(para)

    test_path = "test_report_output.docx"
    doc.save(test_path)
    size = os.path.getsize(test_path)
    record("docx_export", size > 1000, f"Saved to {test_path} ({size} bytes)")
    os.remove(test_path)
except ImportError:
    record("docx_export", False, "python-docx not installed")
except Exception as e:
    record("docx_export", False, str(e))


# ── 5. Chart data export ───────────────────────────────────────
print("\n--- 5. Chart Data Export ---\n")

import pandas as pd

# Build chart data from the analysis results (mocked)
# In real usage this comes from app._analysis_results

mock_classifications = [
    {"text": "Trainerin war toll", "main_category": "Positiv",
     "cluster_title": "Trainer Lob", "cluster_id": "c1", "confidence": "high"},
    {"text": "Zu wenig Pausen", "main_category": "Kritik",
     "cluster_title": "Zeitkritik", "cluster_id": "c2", "confidence": "high"},
    {"text": "Spannende Inhalte", "main_category": "Positiv",
     "cluster_title": "Inhalt positiv", "cluster_id": "c3", "confidence": "medium"},
    {"text": "Trainerin sehr engagiert", "main_category": "Positiv",
     "cluster_title": "Trainer Lob", "cluster_id": "c1", "confidence": "high"},
    {"text": "Zeitplan zu eng", "main_category": "Kritik",
     "cluster_title": "Zeitkritik", "cluster_id": "c2", "confidence": "medium"},
] * 20  # 100 responses


def export_chart_data(question: str, classifications: list, output_dir: str):
    """Export chart data for a single question as formatted Excel.

    Creates two sheets:
      1. Summary: category distribution table (for graphic designer)
      2. Responses: all individual responses with categories
    """
    os.makedirs(output_dir, exist_ok=True)

    # Build summary table
    from collections import Counter
    total = len(classifications)
    cat_counts = Counter()
    for c in classifications:
        key = (c.get("main_category", ""), c.get("cluster_title", ""))
        cat_counts[key] += 1

    summary_rows = []
    for (main, sub), count in cat_counts.most_common():
        summary_rows.append({
            "Main Category": main,
            "Sub-Category": sub,
            "Count": count,
            "Percentage": round(count / total * 100, 1),
        })

    summary_df = pd.DataFrame(summary_rows)

    # Build responses table
    resp_rows = []
    for c in classifications:
        resp_rows.append({
            "Response": c.get("text", ""),
            "Main Category": c.get("main_category", ""),
            "Sub-Category": c.get("cluster_title", ""),
            "Confidence": c.get("confidence", ""),
        })
    resp_df = pd.DataFrame(resp_rows)

    # Clean filename
    import re
    clean_q = re.sub(r'[^\w\s-]', '', question)[:50].strip()
    path = os.path.join(output_dir, f"chart_data_{clean_q}.xlsx")

    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        summary_df.to_excel(writer, sheet_name="Distribution", index=False)
        resp_df.to_excel(writer, sheet_name="All Responses", index=False)

        # Auto-fit column widths
        for sheet_name in writer.sheets:
            ws = writer.sheets[sheet_name]
            for col in ws.columns:
                max_len = max(len(str(cell.value or "")) for cell in col)
                ws.column_dimensions[col[0].column_letter].width = min(max_len + 2, 50)

    return path


# Test export
test_dir = "test_chart_export"
path = export_chart_data(
    "[Post] Was hat Ihnen gut gefallen?",
    mock_classifications, test_dir)

record("chart_data_export",
       os.path.exists(path) and os.path.getsize(path) > 1000,
       f"Exported to {path} ({os.path.getsize(path)} bytes)")

# Verify contents
verify_df = pd.read_excel(path, sheet_name="Distribution")
record("chart_data_summary",
       len(verify_df) >= 2 and "Percentage" in verify_df.columns,
       f"{len(verify_df)} categories, columns: {list(verify_df.columns)}")

verify_resp = pd.read_excel(path, sheet_name="All Responses")
record("chart_data_responses",
       len(verify_resp) == len(mock_classifications),
       f"{len(verify_resp)} responses")

# Show sample
print(f"\n  Summary table:")
print(verify_df.to_string(index=False))

# Cleanup
import shutil
shutil.rmtree(test_dir, ignore_errors=True)


# ── 6. Quantitative data export ────────────────────────────────
print("\n\n--- 6. Quantitative Data Export ---\n")

from bc4d_intel.core.stats_engine import analyze_all_likert, analyze_matched_likert

# Pre-survey stats
pre_items = analyze_all_likert(match_result["pre_all"], pre_roles)
pre_rows = []
for item in pre_items:
    s = item["stats"]
    if s["mean"] is not None:
        pre_rows.append({
            "Item": item["label"][:60],
            "N": s["n"],
            "Mean": s["mean"],
            "SD": s["sd"],
            "CI Lower": s["ci_lower"],
            "CI Upper": s["ci_upper"],
            "% Agree (4-5)": s["pct_agree"],
            "% Disagree (1-2)": s["pct_disagree"],
        })

pre_stats_df = pd.DataFrame(pre_rows)
record("pre_stats_table", len(pre_stats_df) > 0,
       f"{len(pre_stats_df)} Likert items with stats")

# Matched comparison
comparisons = analyze_matched_likert(
    match_result["matched"], pre_roles, post_roles)
comp_rows = []
for c in comparisons:
    comp = c.get("comparison", {})
    if "error" not in comp:
        comp_rows.append({
            "Item": c["label"][:60],
            "Pre Mean": comp.get("pre_mean"),
            "Post Mean": comp.get("post_mean"),
            "Change": comp.get("mean_change"),
            "CI Lower": comp.get("ci_lower"),
            "CI Upper": comp.get("ci_upper"),
            "Cohen's d": comp.get("cohens_d"),
            "Effect": comp.get("effect_label"),
            "p-value": comp.get("p_corrected"),
            "Significant": comp.get("significant"),
            "% Improved": comp.get("improved_pct"),
            "% Declined": comp.get("declined_pct"),
        })

comp_df = pd.DataFrame(comp_rows)
record("matched_stats_table", len(comp_df) > 0,
       f"{len(comp_df)} paired comparisons with effect sizes")

# Export all stats to Excel
test_stats_path = "test_stats_export.xlsx"
with pd.ExcelWriter(test_stats_path, engine="openpyxl") as writer:
    pre_stats_df.to_excel(writer, sheet_name="Pre-Survey", index=False)
    comp_df.to_excel(writer, sheet_name="Pre-Post Comparison", index=False)

    for sheet_name in writer.sheets:
        ws = writer.sheets[sheet_name]
        for col in ws.columns:
            max_len = max(len(str(cell.value or "")) for cell in col)
            ws.column_dimensions[col[0].column_letter].width = min(max_len + 2, 50)

record("stats_export",
       os.path.exists(test_stats_path),
       f"Exported to {test_stats_path}")

print(f"\n  Pre-survey stats ({len(pre_stats_df)} items):")
if len(pre_stats_df) > 0:
    print(pre_stats_df.head(3).to_string(index=False))

print(f"\n  Matched comparisons ({len(comp_df)} items):")
if len(comp_df) > 0:
    print(comp_df.head(3).to_string(index=False))

os.remove(test_stats_path)


# ── Summary ─────────────────────────────────────────────────────
print("\n" + "=" * 70)
print("SUMMARY")
print("=" * 70)

n_pass = sum(1 for r in RESULTS.values() if r["passed"])
n_fail = sum(1 for r in RESULTS.values() if not r["passed"])

for name, r in RESULTS.items():
    print(f"  {'PASS' if r['passed'] else 'FAIL'} -- {name}")

print(f"\n  {n_pass}/{n_pass + n_fail} passed")
sys.exit(0 if n_fail == 0 else 1)
