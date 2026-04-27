"""Build library.docx — BC4D Intel pipeline architecture document."""
import sys
sys.path.insert(0, ".")

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# -- Style setup --
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles["Heading %d" % level]
    h.font.color.rgb = RGBColor(0xC8, 0x17, 0x5D)  # BC4D pink

# ============================================================================
# TITLE
# ============================================================================
title = doc.add_heading("BC4D Intel -- Pipeline Architecture Library", level=0)
title.runs[0].font.color.rgb = RGBColor(0xC8, 0x17, 0x5D)
doc.add_paragraph("Internal reference document -- March 2026")
doc.add_paragraph("")

# ============================================================================
# 1. OVERVIEW
# ============================================================================
doc.add_heading("1. What This App Does", level=1)
doc.add_paragraph(
    "BC4D Intel is a desktop application for evaluating BC4D (Bystander Courage "
    "for Democracy) training programmes run by ISD Deutschland. It processes "
    "Excel-based pre- and post-surveys from each training round (called a "
    "\"Staffel\"), matches participants across both surveys via pseudonymised keys, "
    "analyses Likert-scale and free-text responses, and produces evaluation reports."
)
doc.add_paragraph(
    "The free-text analysis pipeline is the most complex part of the app. It must "
    "categorise hundreds of open-ended German-language responses into meaningful "
    "themes, do so affordably across many Staffels, and run fast enough for a "
    "desktop GUI."
)

# ============================================================================
# 2. PIPELINE ARCHITECTURE
# ============================================================================
doc.add_heading("2. The Free-Text Pipeline -- How It Works", level=1)

doc.add_heading("2.1 High-Level Flow", level=2)
doc.add_paragraph(
    "For each free-text question in a survey (typically 6-9 questions), "
    "the pipeline runs these steps:"
)
steps = [
    ("Cache Check (free, local)",
     "Every new response is compared against a local SQLite database of previously "
     "classified responses using a cross-encoder model. If a sufficiently similar "
     "cached response is found and passes 6 safety checks, the cached classification "
     "is reused. Responses without a good match are sent to the AI pipeline."),
    ("Taxonomy Design (Sonnet, ~$0.025/question)",
     "Claude Sonnet reads ALL uncached responses and designs a hierarchical taxonomy "
     "of themes -- typically 3-5 main categories with 2-3 sub-categories each. Each "
     "category gets a title, example responses, and include/exclude rules."),
    ("Cross-Encoder Classification (free, local)",
     "A local cross-encoder model (ms-marco-MiniLM-L-6-v2) scores every uncached "
     "response against every taxonomy category in one batch. Each response is assigned "
     "to the best-matching category. Confidence is based on the margin between the "
     "top two scores."),
    ("Edge Case Review (Haiku, ~$0.005/question)",
     "Responses classified with low confidence (typically 10-20% of the total) are "
     "sent to Claude Haiku in batches of 20. Haiku picks the correct category from "
     "the existing taxonomy -- a simple closed-list selection task."),
    ("Cache Update",
     "All newly classified responses are added to the SQLite cache so that future "
     "Staffels can reuse them without calling the AI."),
]
for i, (title_text, body) in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run("Step %d: %s" % (i, title_text))
    run.bold = True
    doc.add_paragraph(body)

doc.add_heading("2.2 Model Routing", level=2)
table = doc.add_table(rows=5, cols=4)
table.style = "Light Shading Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Task", "Model", "Cost", "Why This Model"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
rows_data = [
    ["Taxonomy design", "Sonnet", "~$0.025/q", "Needs reasoning about themes in German"],
    ["Classification", "Cross-encoder (local)", "$0", "Fast scoring, no API needed"],
    ["Edge case review", "Haiku", "~$0.005/q", "Simple pick-from-list, 40% cheaper than Sonnet"],
    ["Report writing", "Sonnet", "~$0.10/report", "Quality German prose"],
]
for ri, row_data in enumerate(rows_data):
    for ci, val in enumerate(row_data):
        table.rows[ri + 1].cells[ci].text = val

# ============================================================================
# 3. WHY WE BUILT IT THIS WAY
# ============================================================================
doc.add_heading("3. Why We Built It This Way", level=1)

doc.add_heading("3.1 Problem: API Costs Scale Linearly", level=2)
doc.add_paragraph(
    "The original approach sent every response to Claude for classification. "
    "With ~140 responses per question and 6-9 questions per Staffel, that meant "
    "800-1200 API calls per Staffel. At ~$0.002 per call (Haiku), that is ~$1.60 "
    "per Staffel -- acceptable once, but ISD runs 10-15 Staffels per year. Over "
    "time the same types of responses keep appearing: \"Die Trainerin war toll\", "
    "\"Mehr Praxisbezug\", etc. Paying to re-classify them is waste."
)

doc.add_heading("3.2 Solution: Cache + Hybrid Pipeline", level=2)
doc.add_paragraph(
    "The answer cache stores every classified response in a local SQLite database. "
    "When a new Staffel arrives, the cross-encoder compares new responses against "
    "cached ones. If a match is found (similarity >= 0.78 after sigmoid normalisation), "
    "the classification is reused for free. Only genuinely new responses go to the AI."
)
doc.add_paragraph("Expected cost reduction over 10 Staffels:")
cost_table = doc.add_table(rows=6, cols=4)
cost_table.style = "Light Shading Accent 1"
for i, h in enumerate(["Staffel", "Cache Hit Rate", "AI Calls", "Cost/Staffel"]):
    cost_table.rows[0].cells[i].text = h
cost_rows = [
    ["1 (first run)", "0%", "Full AI", "~$0.65"],
    ["2", "~60%", "~40% AI", "~$0.26"],
    ["3-5", "~80%", "~20% AI", "~$0.13"],
    ["6-10", "~90%", "~10% AI", "~$0.06"],
    ["10+", "~95%", "~5% AI", "~$0.03"],
]
for ri, rd in enumerate(cost_rows):
    for ci, val in enumerate(rd):
        cost_table.rows[ri + 1].cells[ci].text = val

doc.add_heading("3.3 Solution: Cross-Encoder for Speed", level=2)
doc.add_paragraph(
    "The cross-encoder (ms-marco-MiniLM-L-6-v2) replaced per-response API calls "
    "for classification. Instead of calling Claude 135 times to classify 135 "
    "responses, the cross-encoder scores all response-category pairs in a single "
    "batch call locally. This was measured as a 39x speedup (2 seconds vs. 78 seconds) "
    "during the ISD Intel project, where it was first developed."
)

doc.add_heading("3.4 Solution: 6-Layer Safety System", level=2)
doc.add_paragraph(
    "Lowering the cache threshold from 0.90 to 0.78 recovered ~30% more cache hits, "
    "but introduced the risk of false matches. The 6-layer safety system catches "
    "dangerous matches that the similarity score alone would miss:"
)
guards = [
    ("Sentiment Polarity Guard",
     "Blocks matches where one text is clearly positive and the other negative. "
     "Uses keyword-based sentiment detection."),
    ("Negation Asymmetry",
     "Blocks \"nicht informativ\" from matching \"informativ\". "
     "Detects negation words (nicht, kein, nie, kaum, wenig) and requires both "
     "texts to have the same negation status."),
    ("Conditional vs. Definitive",
     "Blocks \"Waere besser gewesen\" from matching \"War gut\". "
     "Detects conditional mood (waere, haette, koennte, sollte) and blocks "
     "if conditional status differs AND sentiments differ."),
    ("Subject Mismatch (Short Responses)",
     "For responses under 6 words, requires at least one content word overlap. "
     "Prevents \"Trainerin toll\" from matching \"Material toll\"."),
    ("Question Context Alignment",
     "On \"Staerken\" questions, blocks negative responses from matching. "
     "On \"Verbesserung\" questions, blocks purely positive responses. "
     "Also checks that the cached category label aligns with the new response's sentiment."),
    ("Short Text Threshold",
     "Responses under 4 words need a similarity score >= 0.95 (vs. 0.78 for longer texts). "
     "Very short texts like \"Gut\" are too ambiguous for reliable matching."),
]
for title_text, body in guards:
    p = doc.add_paragraph()
    run = p.add_run("Guard: %s. " % title_text)
    run.bold = True
    p.add_run(body)

# ============================================================================
# 4. THE PROBLEMS WE ARE FACING
# ============================================================================
doc.add_heading("4. Current Problems", level=1)

doc.add_heading("4.1 CRITICAL: Cross-Encoder is Catastrophically Slow for Cache Matching", level=2)
doc.add_paragraph(
    "The cross-encoder batch approach works well for CLASSIFICATION (responses vs. "
    "~10-20 taxonomy categories = ~2,700 pairs). But for CACHE MATCHING it creates "
    "an O(n * m) pair explosion:"
)
p = doc.add_paragraph()
run = p.add_run(
    "140 responses x 600 cached entries = 84,000 pairs PER QUESTION\n"
    "6 questions per Staffel = 504,000 pairs total"
)
run.bold = True
doc.add_paragraph(
    "In our simulation test, scoring 7,940 pairs (20 sampled responses x 397 cache "
    "entries for a single question) took 442 seconds -- roughly 18 pairs/second. "
    "At that rate, a full Staffel cache check would take approximately 7.8 HOURS."
)
doc.add_paragraph(
    "This makes the app completely unusable. The cache was designed to SAVE time "
    "and money, but the matching step itself has become the bottleneck. The "
    "cross-encoder that gave us a 39x speedup for classification (small matrix) "
    "becomes a 39x slowdown for cache matching (huge matrix)."
)
doc.add_paragraph(
    "The original speed claim of \"2 seconds for 67,500 pairs\" from the embedder "
    "module comments appears to have been measured under different conditions "
    "(possibly GPU, or a much smaller actual pair count). On a typical laptop CPU, "
    "the cross-encoder processes only ~18 pairs/second."
)

doc.add_heading("4.2 CRITICAL: Question Pattern Matching is Brittle", level=2)
doc.add_paragraph(
    "The cache stores simplified question patterns like \"Bitte nennen Sie "
    "Staerken des Kurses.\" but real Excel column names look like \"Bitte nennen "
    "Sie Starken des Kurses. (Angabe erforderlich)\". The normalisation function "
    "only strips [Pre]/[Post] prefixes and truncates to 100 characters -- it does "
    "NOT handle:"
)
bullets = [
    "Umlauts vs. ASCII (Starken vs. Staerken)",
    "Trailing metadata like \"(Angabe erforderlich)\" or \"(Mehrfachnennungen moeglich)\"",
    "Whitespace and newline differences in column headers",
    "Slight rephrasing across different Staffel survey versions",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")
doc.add_paragraph(
    "Result: without a manual column-to-pattern mapping, the cache returns 0% hits "
    "because the question lookup fails before any response matching even begins. "
    "This was confirmed in our simulation where all three datasets showed 0% hits "
    "until we added a manual mapping dictionary."
)

doc.add_heading("4.3 HIGH: Cache Hit Rates Lower Than Expected", level=2)
doc.add_paragraph(
    "Even after fixing the question matching, the positive dataset (which should "
    "represent easy-to-match responses) showed only 35% cache hits on the first "
    "column tested. The target was >70%. Possible causes:"
)
bullets2 = [
    "The synthetic data may use different phrasing than the original Staffel data "
    "that populated the cache.",
    "The 0.78 threshold may still be too conservative for paraphrased responses.",
    "The safety guards may be over-blocking legitimate matches (false negatives).",
    "The cache may not have enough diversity -- 400-800 entries per question may "
    "not cover the range of ways people express similar ideas.",
]
for b in bullets2:
    doc.add_paragraph(b, style="List Bullet")
doc.add_paragraph(
    "Note: we could not complete the full simulation to determine which factor "
    "dominates, because the speed problem (4.1) made it impractical to run."
)

doc.add_heading("4.4 MEDIUM: Simulation Cannot Run End-to-End", level=2)
doc.add_paragraph(
    "The 5-phase simulation plan was designed to test cache hit rates, safety "
    "guards, full pipeline, classification quality, and architecture decisions. "
    "We could only partially complete Phase 1 before the cross-encoder speed "
    "made further testing impractical. Phase 2 (safety guard unit tests) would "
    "run instantly but depends on the cross-encoder for the question-context "
    "tests. Phases 3-5 were never reached."
)

# ============================================================================
# 5. WHAT WORKS WELL
# ============================================================================
doc.add_heading("5. What Works Well", level=1)
works = [
    ("Hybrid taxonomy + classification",
     "The approach of using Sonnet for taxonomy design and the cross-encoder for "
     "classification is sound. For the classification step (responses vs. ~15 "
     "categories), the cross-encoder is fast and accurate."),
    ("6-layer safety system (logic)",
     "The safety guards are well-designed and cover the right edge cases. The "
     "sentiment, negation, conditional, subject, and context checks are the "
     "correct safeguards for German survey data."),
    ("Model routing",
     "Using Sonnet for reasoning and Haiku for simple tasks saves ~40% on edge "
     "case costs with no quality loss."),
    ("Cost model",
     "The projected cost reduction from caching (70% over 10 Staffels) is "
     "compelling and the per-question costs (~$0.03-0.04) are very low."),
    ("State persistence",
     "Checkpoint saves after each question provide crash resilience. The app can "
     "resume from where it left off."),
]
for title_text, body in works:
    p = doc.add_paragraph()
    run = p.add_run("%s. " % title_text)
    run.bold = True
    p.add_run(body)

# ============================================================================
# 6. RECOMMENDED FIXES
# ============================================================================
doc.add_heading("6. Recommended Fixes (Priority Order)", level=1)

doc.add_heading("6.1 Fix Cross-Encoder Speed (CRITICAL)", level=2)
doc.add_paragraph("The O(n*m) pair explosion must be eliminated. Options:")
fixes = [
    ("Pre-filter with TF-IDF",
     "Before cross-encoder scoring, compute TF-IDF vectors for all responses and "
     "cache entries. Use cosine similarity to select the top-50 candidate cache "
     "entries per response. Then run the cross-encoder only on those 50 candidates "
     "instead of all 600. Reduces pairs from 84,000 to 7,000 per question (~12x faster). "
     "Simple to implement, no new dependencies."),
    ("Bi-encoder + FAISS index",
     "Precompute sentence embeddings for all cache entries using a bi-encoder "
     "(e.g., all-MiniLM-L6-v2). Store in a FAISS index. At query time, encode "
     "new responses and retrieve top-k nearest neighbours (~20-50x faster). "
     "More complex but scales to millions of cache entries."),
    ("Cap cache diversity",
     "Limit cache to ~200 diverse entries per question by clustering and keeping "
     "only cluster centroids. Reduces pairs by 3x with minimal coverage loss. "
     "Can combine with option A or B."),
]
for title_text, body in fixes:
    p = doc.add_paragraph()
    run = p.add_run("Option: %s. " % title_text)
    run.bold = True
    p.add_run(body)

doc.add_heading("6.2 Fix Question Pattern Matching (CRITICAL)", level=2)
doc.add_paragraph(
    "Replace exact-match lookup with fuzzy matching. Use difflib.SequenceMatcher "
    "(ratio > 0.6) or extract keywords from the question and match on keyword overlap. "
    "Alternatively, normalise both sides more aggressively: strip parenthetical "
    "metadata, replace umlauts, collapse whitespace."
)

doc.add_heading("6.3 Investigate Cache Hit Rates (HIGH)", level=2)
doc.add_paragraph(
    "Once speed is fixed, re-run the simulation to determine whether the 35% hit "
    "rate is caused by the threshold, the safety guards, or the data. Log which "
    "safety guard blocked each rejected match to identify over-blocking."
)

# ============================================================================
# 7. FILE REFERENCE
# ============================================================================
doc.add_heading("7. File Reference", level=1)
files = [
    ("bc4d_intel/core/answer_cache.py",
     "SQLite cache + cross-encoder matching + 6-layer safety. Key functions: "
     "classify_from_cache(), add_to_cache(), _safe_to_cache_match()"),
    ("bc4d_intel/core/embedder.py",
     "Hybrid pipeline: taxonomy design (Sonnet) + cross-encoder classification + "
     "edge case review (Haiku). Key function: full_pipeline()"),
    ("bc4d_intel/screens/screen_analysis.py",
     "Orchestrator: cache-first, then AI, then merge, then save. Runs in background thread."),
    ("bc4d_intel/ai/claude_client.py",
     "Unified API client with retry logic. Routes to Haiku or Sonnet based on task type."),
    ("bc4d_intel/ai/prompts.py",
     "System prompts for tagging (10 categories) and report writing (7 sections)."),
    ("bc4d_intel/constants.py",
     "Model routing (AI_MODELS), theme colours, Likert scale, navigation."),
    ("bc4d_intel/app_state.py",
     "Central state object. All data flows through AppState. Persists to sessions/latest.bc4d."),
]
file_table = doc.add_table(rows=len(files) + 1, cols=2)
file_table.style = "Light Shading Accent 1"
file_table.rows[0].cells[0].text = "File"
file_table.rows[0].cells[1].text = "Purpose"
for ri, (path, desc) in enumerate(files):
    file_table.rows[ri + 1].cells[0].text = path
    file_table.rows[ri + 1].cells[1].text = desc

# ============================================================================
# SAVE
# ============================================================================
doc.save("library.docx")
print("Saved: library.docx")
